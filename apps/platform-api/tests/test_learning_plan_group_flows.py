from __future__ import annotations

import base64
import io
import hashlib
import json
import sys
from pathlib import Path

import pytest
from docx import Document


REPO_ROOT = Path(__file__).resolve().parents[3]
SCRIPTS_ROOT = REPO_ROOT / "scripts"
sys.path.insert(0, str(SCRIPTS_ROOT))

from apply_group_meeting_flow_adjustments import build_candidate_catalog  # noqa: E402
from parse_group_meeting_flows import (  # noqa: E402
    _learning_content_type,
    _learning_content_titles,
    load_course_rules,
    parse_flow,
)
from app.api.learning_plans import (  # noqa: E402
    GROUP_MEETING_CREDIT_POLICY,
    _learning_plan_group_meeting_flow_payload,
)


INVENTORY = REPO_ROOT / "data/learning-plans/group-meeting-source-inventory-2026.json"
FLOWS = REPO_ROOT / "data/learning-plans/group-meeting-flows-2026.1.json"
MAPPING = REPO_ROOT / "data/learning-plans/cycle-flow-mapping-2026.1.json"
RULES = REPO_ROOT / "data/learning-plans/course-credit-rules-2026.json"
PLAN = REPO_ROOT / "data/learning-plans/standard-3y-2026.json"
REVIEW = REPO_ROOT / "data/learning-plans/standard-3y-2026.review.json"


_ONE_PIXEL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def _sample_docx(tmp_path: Path) -> tuple[Path, Path]:
    source_root = tmp_path / "source"
    source_root.mkdir()
    path = source_root / "【2026版】第1学年" / "【第1个月】2026版1、4、7、10月开班.docx"
    path.parent.mkdir()
    document = Document()
    document.add_paragraph("第1个月小组学习会流程")
    for text, with_image in (
        ("1. 观看《如何制作核算表》视频并进行整体核算表研讨", True),
        ("2. 幸福测评表结果检视与改善", True),
        ("3. 班级学习会发表稿编写讲解", True),
        ("4. 空巴", True),
    ):
        paragraph = document.add_paragraph(text)
        if with_image:
            paragraph.add_run().add_picture(io.BytesIO(_ONE_PIXEL_PNG), width=None)
    document.add_paragraph("空巴后班会二维码（应排除）")
    document.save(path)
    return path, source_root


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def test_flow_parser_stops_at_first_konpa(tmp_path: Path) -> None:
    path, source_root = _sample_docx(tmp_path)
    flow = parse_flow(path, source_root, load_course_rules(RULES))

    assert flow["status"] == "PARSED"
    assert flow["boundary"]["terminal_step_count"] == 1
    assert flow["steps"][-1]["is_terminal"] is True
    assert all("空巴" in step["content"] or "空吧" in step["content"] for step in flow["steps"][-1:])


def test_qr_after_konpa_is_excluded(tmp_path: Path) -> None:
    path, source_root = _sample_docx(tmp_path)
    flow = parse_flow(path, source_root, load_course_rules(RULES))

    assert flow["boundary"]["qr_after_terminal_excluded"] is True
    assert all(
        node["source_paragraph_index"] <= flow["steps"][-1]["source_paragraph_index"]
        for node in flow["course_nodes"]
    )


def test_learning_content_is_semantic_and_deduplicates_book_mentions() -> None:
    text = "观看《经营十二条实践》视频，进行《经营十二条》读书和《经营十二条实践》视频学习总结"

    assert _learning_content_titles(text) == ["经营十二条实践"]


def test_learning_content_can_contain_two_videos_without_qr() -> None:
    text = "【视频A】视频学习+研讨；【视频B】视频学习+研讨"

    assert _learning_content_titles(text) == ["视频A", "视频B"]


def test_learning_content_title_list_keeps_titles_before_shared_marker() -> None:
    text = "【改善创新委讲解】、【改善创新优秀企业践行分享】视频学习+研讨；"

    assert _learning_content_titles(text) == [
        "改善创新委讲解",
        "改善创新优秀企业践行分享",
    ]


def test_conditional_video_note_is_not_a_learning_content() -> None:
    assert _learning_content_type(
        "如本月小组学习会中经营分析会实操和其他学习内容（如视频学习）结合，则按全天设计"
    ) is None


def test_learning_content_does_not_inherit_credit_from_another_title() -> None:
    text = "【会计七原则实践】视频与《经营与会计》书籍的学习总结"

    # The title extractor is the boundary used before credit matching.  A
    # book mentioned beside a mapped video must not become another mapped
    # learning content node.
    assert _learning_content_titles(text) == ["会计七原则实践"]


def test_cycle_five_video_is_visible_without_qr() -> None:
    payload = json.loads(FLOWS.read_text(encoding="utf-8"))
    nodes = [
        node
        for flow in payload["flows"]
        if flow["year_index"] == 1 and flow["cycle_index"] == 5
        for node in flow["learning_content_nodes"]
    ]

    assert len(nodes) == 1
    assert nodes[0]["title"] == "关于核算表分析&任务单的制作"
    assert nodes[0]["task_type"] == "VIDEO_LEARNING"
    assert nodes[0]["is_required"] is True
    assert nodes[0]["qr_refs"] == []
    assert nodes[0]["credit_rule_key"] == "Y1-ACCOUNTING-ANALYSIS-TASK"


def test_all_learning_content_nodes_are_not_defined_by_qr_count() -> None:
    payload = json.loads(FLOWS.read_text(encoding="utf-8"))
    content_nodes = [
        node
        for flow in payload["flows"]
        for node in flow["learning_content_nodes"]
    ]

    assert len(content_nodes) == payload["quality_report"]["learning_content_node_count"]
    assert any(node["qr_refs"] == [] for node in content_nodes)
    assert payload["quality_report"]["required_video_without_qr_count"] > 0


def test_first_group_meeting_has_three_course_nodes(tmp_path: Path) -> None:
    path, source_root = _sample_docx(tmp_path)
    flow = parse_flow(path, source_root, load_course_rules(RULES))

    assert len(flow["course_nodes"]) == 3
    assert [node["course_key"] for node in flow["course_nodes"]] == [
        "Y1-INTEGRATED-ACCOUNTING",
        "Y1-HAPPINESS-ASSESSMENT",
        "Y1-CLASS-SPEECH-DRAFT",
    ]


def test_first_year_credit_rules_match_confirmed_values() -> None:
    payload = json.loads(RULES.read_text(encoding="utf-8"))
    points = {rule["course_key"]: rule["credit_points"] for rule in payload["rules"]}

    assert points["Y1-HAPPINESS-ASSESSMENT"] == 20
    assert points["Y1-CLASS-SPEECH-DRAFT"] == 20
    assert points["Y1-SIX-DILIGENCES"] == 20
    assert points["Y1-TWELVE-MANAGEMENT"] == 20
    assert points["Y1-INTEGRATED-ACCOUNTING"] == 40
    assert points["Y1-ACCOUNTING-ANALYSIS-TASK"] == 40
    assert points["Y1-SEVEN-ACCOUNTING-PRINCIPLES"] == 20


def test_kyocera_annual_plan_credit_follows_business_reference() -> None:
    payload = json.loads(RULES.read_text(encoding="utf-8"))
    rule = next(item for item in payload["rules"] if item["course_key"] == "Y1-KYOCERA-ANNUAL-PLAN")
    assert rule["credit_points"] == 40


def test_annual_monthly_management_credit_follows_business_reference() -> None:
    payload = json.loads(RULES.read_text(encoding="utf-8"))
    rule = next(item for item in payload["rules"] if item["course_key"] == "Y1-ANNUAL-MONTHLY-MGMT")
    assert rule["credit_points"] == 11


def test_group_meeting_base_credit_is_cycle_level_not_step_level() -> None:
    plan = json.loads(PLAN.read_text(encoding="utf-8"))
    assert GROUP_MEETING_CREDIT_POLICY["credit_points_per_person"] == 4
    assert GROUP_MEETING_CREDIT_POLICY["task_level_credit_points"] is None
    assert all(
        task["credit_points"] is None
        for track in plan["cohort_tracks"]
        for cycle in track["cycles"]
        for task in cycle["tasks"]
        if task["task_type"] == "GROUP_MEETING"
    )


def test_unknown_course_credit_stays_null() -> None:
    payload = json.loads(FLOWS.read_text(encoding="utf-8"))
    unknown = [
        node
        for flow in payload["flows"]
        for node in flow["course_nodes"]
        if node["credit_status"] == "QR_REVIEW_REQUIRED"
    ]
    assert unknown
    assert all(node["credit_points"] is None for node in unknown)


def test_success_formula_49_day_video_has_no_course_credit() -> None:
    payload = json.loads(FLOWS.read_text(encoding="utf-8"))
    nodes = [
        node
        for flow in payload["flows"]
        for node in flow["learning_content_nodes"]
        if node["title"] == "成功方程式49天讲解"
    ]
    assert nodes
    assert all(node["credit_rule_key"] is None for node in nodes)
    assert all(node["credit_points"] is None for node in nodes)


def test_confirmed_screenshot_flows_have_only_cycle_attendance_credit() -> None:
    payload = json.loads(FLOWS.read_text(encoding="utf-8"))
    flows = {
        flow["flow_key"]: flow
        for flow in payload["flows"]
        if flow.get("status") == "MANUALLY_CONFIRMED"
    }
    assert set(flows) == {
        "Y3-C28-COHORT-1-4-BIZCONF",
        "Y3-C29-COHORT-1-4-BIZCONF",
        "Y3-C32-COHORT-7-BIZCONF",
        "Y3-C33-COHORT-7-BIZCONF",
        "Y3-C34-COHORT-7-BIZCONF",
    }
    assert all(flow["learning_content_nodes"] == [] for flow in flows.values())
    assert all(flow["course_nodes"] == [] for flow in flows.values())
    assert all(
        all(step["qr_refs"] == [] for step in flow["steps"])
        for flow in flows.values()
    )
    assert all(flow["business_confirmation"]["attendance_credit_points"] == 4 for flow in flows.values())


def test_cycle_flow_mapping_does_not_use_calendar_month_as_primary_key() -> None:
    payload = json.loads(MAPPING.read_text(encoding="utf-8"))
    assert payload["quality_report"]["entry_count"] == 144
    assert payload["quality_report"]["template_count"] == 4
    assert payload["quality_report"]["template_cycle_definition"] == "4个开班月份模板 × 36学习周期"
    assert payload["quality_report"]["calendar_month_used_as_primary_key"] is False
    assert all(not item["calendar_month_used_as_primary_key"] for item in payload["mappings"])
    sample = next(
        item
        for item in payload["mappings"]
        if item["cohort_month"] == 1 and item["learning_cycle_index"] == 26
    )
    assert sample["template_label"] == "1月开班模板"
    assert sample["learning_cycle_label"] == "1月开班模板 · 第26学习周期"
    assert sample["lookup_key"]["learning_cycle_index"] == 26


def test_cycle_26_duplicate_selection_and_screenshot_mappings_are_active() -> None:
    mapping = json.loads(MAPPING.read_text(encoding="utf-8"))
    by_key = {item["mapping_key"]: item for item in mapping["mappings"]}
    for mapping_key in ("1-26", "4-26", "7-26"):
        assert by_key[mapping_key]["status"] == "MAPPED"
        assert by_key[mapping_key]["flow_key"] == "Y3-C26-COHORT-1-4-7-c5846a2908"
    for mapping_key in ("1-28", "4-28", "1-29", "4-29", "7-32", "7-33", "7-34"):
        assert by_key[mapping_key]["status"] == "MAPPED"
        assert by_key[mapping_key]["flow_key"].endswith("BIZCONF")

    flows = json.loads(FLOWS.read_text(encoding="utf-8"))["flows"]
    duplicate = next(
        flow for flow in flows if flow["flow_key"] == "Y3-C26-COHORT-1-4-7-3131994926"
    )
    assert duplicate["status"] == "DUPLICATE_SUPERSEDED"
    assert duplicate["superseded_by"] == "Y3-C26-COHORT-1-4-7-c5846a2908"


def _valid_adjustment() -> tuple[dict, dict]:
    base_catalog = json.loads(FLOWS.read_text(encoding="utf-8"))
    review = json.loads(REVIEW.read_text(encoding="utf-8"))
    inventory = json.loads(INVENTORY.read_text(encoding="utf-8"))
    flow = base_catalog["flows"][0]
    adjustment = {
        "status": "DRAFT",
        "candidate_version_label": "2026.1",
        "overwrite_confirmed": False,
        "base_source_commit": review["source_commit"],
        "base_source_json_sha256": review["source_json_sha256"],
        "base_source_workbooks": review["source_workbooks"],
        "base_group_flow_source_files": [
            {key: item[key] for key in ("filename", "relative_path", "sha256")}
            for item in inventory["included_files"]
        ],
        "base_course_credit_rules_sha256": _sha256(RULES),
        "changes": [{
            "flow_key": flow["flow_key"],
            "steps": [{
                "title": step["title"],
                "content": step["content"],
                "is_required": step["is_required"],
                "notes": None,
            } for step in flow["steps"]],
            "notes": "测试调整",
        }],
    }
    return base_catalog, adjustment


def test_adjustment_export_binds_all_fingerprints() -> None:
    base_catalog, adjustment = _valid_adjustment()
    review = json.loads(REVIEW.read_text(encoding="utf-8"))
    candidate = build_candidate_catalog(
        base_catalog,
        adjustment,
        review_manifest=review,
        base_json=PLAN,
        rules_json=RULES,
        inventory_json=INVENTORY,
    )
    lineage = candidate["source"]
    assert lineage["base_source_commit"] == review["source_commit"]
    assert lineage["base_source_json_sha256"] == review["source_json_sha256"]
    assert lineage["base_source_workbooks"] == review["source_workbooks"]
    assert lineage["base_group_flow_source_files"]
    assert lineage["base_course_credit_rules_sha256"] == _sha256(RULES)


def test_confirmed_2026_cannot_be_overwritten() -> None:
    base_catalog, adjustment = _valid_adjustment()
    adjustment["overwrite_confirmed"] = True
    with pytest.raises(ValueError, match="覆盖"):
        build_candidate_catalog(
            base_catalog,
            adjustment,
            review_manifest=json.loads(REVIEW.read_text(encoding="utf-8")),
            base_json=PLAN,
            rules_json=RULES,
            inventory_json=INVENTORY,
        )


def test_group_flow_catalog_contains_read_only_source_evidence() -> None:
    payload = _learning_plan_group_meeting_flow_payload()
    assert payload["source_fragment_count"] == 575
    assert payload["flow_count"] == 83
    assert len(payload["base_group_flow_source_files"]) == 78
    assert payload["base_course_credit_rules_sha256"] == _sha256(RULES)
    assert payload["quality_report"]["manually_confirmed_flow_count"] == 5
    assert payload["quality_report"]["superseded_flow_count"] == 1
